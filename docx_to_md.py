#!/usr/bin/env python3
"""Simple .docx to .md converter for SA_Toolbelt documentation."""
import sys
from docx import Document

def convert_docx_to_md(docx_path, md_path=None):
    if md_path is None:
        md_path = docx_path.rsplit('.', 1)[0] + '.md'

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style else ''
        text = para.text.strip()

        if not text and not style.startswith('heading'):
            lines.append('')
            continue

        # Map Word heading styles to markdown
        if 'heading 1' in style:
            lines.append(f'# {text}')
        elif 'heading 2' in style:
            lines.append(f'## {text}')
        elif 'heading 3' in style:
            lines.append(f'## {text}')
        elif 'heading 4' in style:
            lines.append(f'### {text}')
        elif 'list' in style or style.startswith('list'):
            lines.append(f'- {text}')
        else:
            # Check for bold/italic runs
            parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    parts.append(f'***{t}***')
                elif run.bold:
                    parts.append(f'**{t}**')
                elif run.italic:
                    parts.append(f'*{t}*')
                else:
                    parts.append(t)
            line = ''.join(parts) if parts else text
            lines.append(line)

        lines.append('')

    # Handle tables
    for table in doc.tables:
        lines.append('')
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append('| ' + ' | '.join(headers) + ' |')
        lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')
        lines.append('')

    # Clean up excessive blank lines
    output = '\n'.join(lines)
    while '\n\n\n' in output:
        output = output.replace('\n\n\n', '\n\n')

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(output.strip() + '\n')

    print(f"Converted: {docx_path} -> {md_path}")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 docx_to_md.py <input.docx> [output.md]")
        sys.exit(1)
    docx_path = sys.argv[1]
    md_path = sys.argv[2] if len(sys.argv) > 2 else None
    convert_docx_to_md(docx_path, md_path)
